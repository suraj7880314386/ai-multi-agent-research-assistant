"""Document ingestion, parsing, and chunking pipeline."""

import os
import uuid
import logging
from datetime import datetime
from typing import List, Dict, Tuple

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document

from app.config import settings

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Handles document upload, parsing, and chunking."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        self.upload_dir = "./uploads"
        os.makedirs(self.upload_dir, exist_ok=True)

        # Track indexed documents
        self._documents: Dict[str, dict] = {}

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from a plain text or markdown file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def extract_text(self, file_path: str) -> str:
        """Route to the appropriate text extractor based on file extension."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self._extract_text_from_docx(file_path)
        elif ext in {".txt", ".md"}:
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def process_document(
        self, file_path: str, filename: str
    ) -> Tuple[str, List[Document]]:
        """
        Full pipeline: extract text → chunk → create Document objects with metadata.

        Returns:
            Tuple of (doc_id, list of LangChain Document chunks)
        """
        doc_id = str(uuid.uuid4())[:8]

        logger.info(f"Processing document: {filename} (id={doc_id})")

        # Extract raw text
        raw_text = self.extract_text(file_path)

        if not raw_text.strip():
            raise ValueError(f"No text could be extracted from {filename}")

        # Split into chunks
        chunks = self.text_splitter.split_text(raw_text)

        # Create LangChain Document objects with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "doc_id": doc_id,
                    "filename": filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "source": filename,
                },
            )
            documents.append(doc)

        # Track document
        self._documents[doc_id] = {
            "doc_id": doc_id,
            "filename": filename,
            "num_chunks": len(chunks),
            "status": "indexed",
            "created_at": datetime.utcnow(),
        }

        logger.info(f"Document {filename} processed: {len(chunks)} chunks created")
        return doc_id, documents

    async def save_upload(self, file) -> str:
        """Save an uploaded file to disk and return the file path."""
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        file_path = os.path.join(self.upload_dir, f"{uuid.uuid4()}{ext}")
        content = await file.read()

        # Check file size
        size_mb = len(content) / (1024 * 1024)
        if size_mb > settings.max_upload_size_mb:
            raise ValueError(
                f"File too large: {size_mb:.1f}MB (max: {settings.max_upload_size_mb}MB)"
            )

        with open(file_path, "wb") as f:
            f.write(content)

        return file_path

    def get_document(self, doc_id: str) -> dict:
        """Get document metadata by ID."""
        if doc_id not in self._documents:
            raise KeyError(f"Document {doc_id} not found")
        return self._documents[doc_id]

    def list_documents(self) -> List[dict]:
        """List all indexed documents."""
        return list(self._documents.values())

    def remove_document(self, doc_id: str) -> bool:
        """Remove document from tracking."""
        if doc_id in self._documents:
            del self._documents[doc_id]
            return True
        return False


# Singleton
document_loader = DocumentLoader()
